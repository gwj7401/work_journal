import io
import urllib.parse
from datetime import date
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from database import get_db
from models.aggregate import QuarterlySummary, AnnualSummary
from models.summary import MonthlySummary
from models.history import SummaryHistory
from models.journal import JournalEntry
from models.user import User
from schemas import (
    QuarterlySummaryGenerate, QuarterlySummaryUpdate, QuarterlySummaryOut,
    AnnualSummaryGenerate, AnnualSummaryUpdate, AnnualSummaryOut,
)
from auth import get_current_user
from services.aggregate_service import (
    generate_quarterly_summary, generate_annual_summary, QUARTER_MONTHS, QUARTER_NAMES
)
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

router = APIRouter(prefix="/api/aggregate", tags=["聚合总结"])

# ─────────────────── 工具 ───────────────────

import os
from docxtpl import DocxTemplate

def _word_export(content: str, filename: str, year: int, render_period: str, user_name: str, dept_name: str):
    """当用户没有挂载模板时的强力兜底引擎：采用系统级别的《平时考核登记表》画框"""
    fallback_template = os.path.join(os.path.dirname(os.path.dirname(__file__)), "sys_table_template.docx")
    doc = DocxTemplate(fallback_template)
    context = {
        'content': content,
        'year': year,
        'render_period': render_period,
        'user_name': user_name,
        'dept_name': dept_name
    }
    doc.render(context)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    encoded_filename = urllib.parse.quote(filename)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"},
    )


def _get_monthly_summaries_for_months(db, user_id, year, months):
    rows = db.query(MonthlySummary).filter(
        MonthlySummary.user_id == user_id,
        MonthlySummary.year == year,
        MonthlySummary.month.in_(months),
    ).all()
    return [
        {"month": r.month, "content": r.edited_content or r.ai_content or ""}
        for r in rows if (r.edited_content or r.ai_content)
    ]


def _get_journal_entries_for_months(db, user_id, year, months):
    start = date(year, min(months), 1)
    end_m = max(months)
    end = date(year + 1, 1, 1) if end_m == 12 else date(year, end_m + 1, 1)
    rows = db.query(JournalEntry).filter(
        JournalEntry.user_id == user_id,
        JournalEntry.entry_date >= start,
        JournalEntry.entry_date < end,
    ).all()
    return [
        {"date": r.entry_date.isoformat(), "content": r.content, "tags": r.tags or []}
        for r in rows if int(r.entry_date.strftime("%m")) in months
    ]


# ─────────────────── 季度总结 ───────────────────

@router.post("/quarterly/generate", response_model=QuarterlySummaryOut, summary="AI生成季度总结")
async def generate_quarterly(
    data: QuarterlySummaryGenerate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if data.quarter not in (1, 2, 3, 4):
        raise HTTPException(400, "季度必须为1-4")
    months = QUARTER_MONTHS[data.quarter]
    monthly_sums = _get_monthly_summaries_for_months(db, current_user.id, data.year, months)

    obj = db.query(QuarterlySummary).filter(
        QuarterlySummary.user_id == current_user.id,
        QuarterlySummary.year == data.year,
        QuarterlySummary.quarter == data.quarter,
    ).first()
    if obj and obj.is_final:
        raise HTTPException(status_code=400, detail="该总结已标记为正式版，如需重新生成请先取消锁定。")

    try:
        ai_text = await generate_quarterly_summary(
            year=data.year, quarter=data.quarter,
            monthly_summaries=monthly_sums,
            dept_name=current_user.dept,
        )
    except RuntimeError as e:
        raise HTTPException(503, str(e))

    if obj:
        obj.ai_content = ai_text
        if not obj.is_final:
            obj.edited_content = ai_text
    else:
        obj = QuarterlySummary(user_id=current_user.id, year=data.year,
                               quarter=data.quarter, ai_content=ai_text, edited_content=ai_text)
        db.add(obj)
    db.commit(); db.refresh(obj)
    return obj


@router.get("/quarterly/{year}/{quarter}", response_model=QuarterlySummaryOut, summary="获取季度总结")
def get_quarterly(year: int, quarter: int, db: Session = Depends(get_db),
                  current_user: User = Depends(get_current_user)):
    obj = db.query(QuarterlySummary).filter(
        QuarterlySummary.user_id == current_user.id,
        QuarterlySummary.year == year, QuarterlySummary.quarter == quarter,
    ).first()
    if not obj:
        raise HTTPException(404, "总结不存在，请先生成")
    return obj


@router.put("/quarterly/{year}/{quarter}", response_model=QuarterlySummaryOut, summary="编辑季度总结")
def update_quarterly(year: int, quarter: int, data: QuarterlySummaryUpdate,
                     db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    obj = db.query(QuarterlySummary).filter(
        QuarterlySummary.user_id == current_user.id,
        QuarterlySummary.year == year, QuarterlySummary.quarter == quarter,
    ).first()
    if not obj:
        raise HTTPException(404, "总结不存在")
    if data.edited_content is not None:
        obj.edited_content = data.edited_content
    if data.is_final is not None:
        if data.is_final and not obj.is_final:
            db.add(SummaryHistory(
                user_id=current_user.id, summary_type="quarterly", 
                summary_id=obj.id, year=obj.year, quarter=obj.quarter,
                content=obj.edited_content or obj.ai_content, version_note="季度定稿存档"
            ))
        obj.is_final = data.is_final
    db.commit(); db.refresh(obj); return obj


@router.post("/quarterly/{year}/{quarter}/upload", response_model=QuarterlySummaryOut, summary="上传Word覆盖季度总结")
async def upload_quarterly_word(
    year: int, quarter: int, file: UploadFile = File(...),
    db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
):
    if not file.filename.endswith((".docx", ".doc")): raise HTTPException(400, "仅支持 .docx 文件")
    obj = db.query(QuarterlySummary).filter(
        QuarterlySummary.user_id == current_user.id,
        QuarterlySummary.year == year, QuarterlySummary.quarter == quarter,
    ).first()
    if not obj: raise HTTPException(404, "总结不存在")
    try:
        content = await file.read()
        doc = Document(io.BytesIO(content))
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        if text:
            obj.edited_content = text
            obj.is_final = True
            db.add(SummaryHistory(
                user_id=current_user.id, summary_type="quarterly", 
                summary_id=obj.id, year=obj.year, quarter=obj.quarter,
                content=text, version_note="季度Word上传存档"
            ))
            db.commit(); db.refresh(obj)
        else: raise HTTPException(400, "未能从文档中提取到文字内容")
    except Exception as e: raise HTTPException(400, f"解析Word文档失败: {str(e)}")
    return obj


@router.get("/quarterly/{year}/{quarter}/export", summary="导出季度总结Word")
def export_quarterly(year: int, quarter: int, db: Session = Depends(get_db),
                     current_user: User = Depends(get_current_user)):
    obj = db.query(QuarterlySummary).filter(
        QuarterlySummary.user_id == current_user.id,
        QuarterlySummary.year == year, QuarterlySummary.quarter == quarter,
    ).first()
    if not obj:
        raise HTTPException(404, "总结不存在")
    content = obj.edited_content or obj.ai_content or ""
    
    # 执行过滤与 RichText 换行构建
    from docxtpl import RichText
    rt = RichText()
    lines = content.split('\n')
    valid_lines = []
    for lb in lines:
        lb = lb.strip()
        if not lb: continue
        lb = lb.replace('**', '').replace('#', '').strip()
        
        # 清理无序列表符号（-、*、•等）
        if lb.startswith("-") or lb.startswith("*") or lb.startswith("•"):
            lb = lb.lstrip("-*•").strip()
            
        # 清除AI常自带的前方多余抬头总结
        if ("工作总结" in lb and (str(year) in lb or "三十" in lb)) and len(lb) < 40:
            continue
        # 去除原生 AI 生成的结语客套话
        if "以上是" in lb and "请领导审阅" in lb:
            continue
        valid_lines.append(lb)

    # 补入用户强烈要求的个人党建及修养心得专属结语
    valid_lines.append("作为一名分管办公室、综合检验检测站的副院长，我不仅要不断加强政治理论知识、专业知识和领导艺术的学习，还要加强党性修养，做廉洁自律的表率，立足本职岗位，防微杜渐，做忠诚、干净、担当，让党和人民放心的干部。")

    for i, lb in enumerate(valid_lines):
        rt.add("　　" + lb)
        if i < len(valid_lines) - 1:
            rt.add('\n')
            
    template_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "docs", "模板", "季度总结模板.docx")
    filename = f"{year}年{QUARTER_NAMES[quarter]}工作总结.docx"
    if os.path.exists(template_path):
        doc = DocxTemplate(template_path)
        context = {
            'content': rt, 'year': year, 'quarter': QUARTER_NAMES[quarter],
            'dept_name': current_user.dept if current_user.dept else "",
            'user_name': current_user.display_name if getattr(current_user, 'display_name', None) else ""
        }
        doc.render(context)
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        encoded_filename = urllib.parse.quote(filename)
        return StreamingResponse(
            buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"}
        )
        
    return _word_export(content, filename, year, QUARTER_NAMES[quarter], current_user.display_name if getattr(current_user, 'display_name', None) else "", current_user.dept if current_user.dept else "")


# ─────────────────── 年度总结 ───────────────────

@router.post("/annual/generate", response_model=AnnualSummaryOut, summary="AI生成年度总结")
async def generate_annual(
    data: AnnualSummaryGenerate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 优先取季度总结
    qrows = db.query(QuarterlySummary).filter(
        QuarterlySummary.user_id == current_user.id,
        QuarterlySummary.year == data.year,
    ).all()
    q_sums = [{"quarter": r.quarter, "content": r.edited_content or r.ai_content or ""} for r in qrows
              if (r.edited_content or r.ai_content)]

    # 次优：月度总结
    m_sums = _get_monthly_summaries_for_months(db, current_user.id, data.year, list(range(1, 13)))

    obj = db.query(AnnualSummary).filter(
        AnnualSummary.user_id == current_user.id,
        AnnualSummary.year == data.year,
    ).first()
    if obj and obj.is_final:
        raise HTTPException(status_code=400, detail="该总结已标记为正式版，如需重新生成请先取消锁定。")

    try:
        ai_text = await generate_annual_summary(
            year=data.year,
            quarterly_summaries=q_sums,
            monthly_summaries=m_sums,
            dept_name=current_user.dept,
        )
    except RuntimeError as e:
        raise HTTPException(503, str(e))

    if obj:
        obj.ai_content = ai_text
        if not obj.is_final:
            obj.edited_content = ai_text
    else:
        obj = AnnualSummary(user_id=current_user.id, year=data.year,
                            ai_content=ai_text, edited_content=ai_text)
        db.add(obj)
    db.commit(); db.refresh(obj); return obj


@router.get("/annual/{year}", response_model=AnnualSummaryOut, summary="获取年度总结")
def get_annual(year: int, db: Session = Depends(get_db),
               current_user: User = Depends(get_current_user)):
    obj = db.query(AnnualSummary).filter(
        AnnualSummary.user_id == current_user.id, AnnualSummary.year == year,
    ).first()
    if not obj:
        raise HTTPException(404, "总结不存在，请先生成")
    return obj


@router.put("/annual/{year}", response_model=AnnualSummaryOut, summary="编辑年度总结")
def update_annual(year: int, data: AnnualSummaryUpdate,
                  db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    obj = db.query(AnnualSummary).filter(
        AnnualSummary.user_id == current_user.id, AnnualSummary.year == year,
    ).first()
    if not obj:
        raise HTTPException(404, "总结不存在")
    if data.edited_content is not None:
        obj.edited_content = data.edited_content
    if data.is_final is not None:
        if data.is_final and not obj.is_final:
            db.add(SummaryHistory(
                user_id=current_user.id, summary_type="annual", 
                summary_id=obj.id, year=obj.year,
                content=obj.edited_content or obj.ai_content, version_note="年度定稿存档"
            ))
        obj.is_final = data.is_final
    db.commit(); db.refresh(obj); return obj


@router.post("/annual/{year}/upload", response_model=AnnualSummaryOut, summary="上传Word覆盖年度总结")
async def upload_annual_word(
    year: int, file: UploadFile = File(...),
    db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
):
    if not file.filename.endswith((".docx", ".doc")): raise HTTPException(400, "仅支持 .docx 文件")
    obj = db.query(AnnualSummary).filter(
        AnnualSummary.user_id == current_user.id, AnnualSummary.year == year,
    ).first()
    if not obj: raise HTTPException(404, "总结不存在")
    try:
        content = await file.read()
        doc = Document(io.BytesIO(content))
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        if text:
            obj.edited_content = text
            obj.is_final = True
            db.add(SummaryHistory(
                user_id=current_user.id, summary_type="annual", 
                summary_id=obj.id, year=obj.year,
                content=text, version_note="年度Word上传存档"
            ))
            db.commit(); db.refresh(obj)
        else: raise HTTPException(400, "未能从文档中提取到文字内容")
    except Exception as e: raise HTTPException(400, f"解析Word文档失败: {str(e)}")
    return obj


@router.get("/annual/{year}/export", summary="导出年度总结Word")
def export_annual(year: int, db: Session = Depends(get_db),
                  current_user: User = Depends(get_current_user)):
    obj = db.query(AnnualSummary).filter(
        AnnualSummary.user_id == current_user.id, AnnualSummary.year == year,
    ).first()
    if not obj:
        raise HTTPException(404, "总结不存在")
    content = obj.edited_content or obj.ai_content or ""
    
    # 执行过滤与 RichText 换行构建
    from docxtpl import RichText
    rt = RichText()
    lines = content.split('\n')
    valid_lines = []
    for lb in lines:
        lb = lb.strip()
        if not lb: continue
        lb = lb.replace('**', '').replace('#', '').strip()
        
        # 清理无序列表符号（-、*、•等）
        if lb.startswith("-") or lb.startswith("*") or lb.startswith("•"):
            lb = lb.lstrip("-*•").strip()
            
        # 清除AI常自带的前方多余抬头总结
        if ("工作总结" in lb and (str(year) in lb or "三十" in lb)) and len(lb) < 40:
            continue
        # 去除原生 AI 生成的结语客套话
        if "以上是" in lb and "请领导审阅" in lb:
            continue
        valid_lines.append(lb)

    # 补入用户强烈要求的个人党建及修养心得专属结语
    valid_lines.append("作为一名分管仪器设备、信息化的质量技术部副部长，我不仅要不断加强政治理论知识、专业知识和领导艺术的学习，还要加强党性修养，做廉洁自律的表率，立足本职岗位，防微杜渐，做忠诚、干净、担当，让党和人民放心的干部。")

    for i, lb in enumerate(valid_lines):
        rt.add("　　" + lb)
        if i < len(valid_lines) - 1:
            rt.add('\n')
            
    template_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "docs", "模板", "年度总结模板.docx")
    filename = f"{year}年度工作总结.docx"
    if os.path.exists(template_path):
        doc = DocxTemplate(template_path)
        context = {
            'content': rt, 'year': year, 
            'dept_name': current_user.dept if current_user.dept else "",
            'user_name': current_user.display_name if getattr(current_user, 'display_name', None) else ""
        }
        doc.render(context)
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        encoded_filename = urllib.parse.quote(filename)
        return StreamingResponse(
            buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"}
        )
        
    return _word_export(content, filename, year, "年度总结", current_user.display_name if getattr(current_user, 'display_name', None) else "", current_user.dept if current_user.dept else "")
