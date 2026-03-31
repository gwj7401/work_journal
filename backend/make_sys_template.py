import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_table_template():
    doc = Document()
    # 页面设置 A4 和公文边距
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.7); sec.bottom_margin = Cm(3.5)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.6)

    # 主标题
    p_title = doc.add_paragraph()
    p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("事业单位工作人员考核登记表")
    run_title.font.name = "方正小标宋简体"
    run_title.font.size = Pt(22)  # 二号子
    rpr = run_title._r.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), '方正小标宋简体')
    rpr.insert(0, rf)

    # 副标题（月份/季度）
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("（{{ year }}年 {{ render_period }}）")
    run_sub.font.name = "楷体_GB2312"
    run_sub.font.size = Pt(16)
    rpr = run_sub._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), '楷体_GB2312'); rpr.insert(0, rf)

    doc.add_paragraph()  # 留白

    # 创建登记表格 
    # 行：姓名/单位、分管、摘要小结、各种意见
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    
    # 所有的表行高设定和默认居中可以在写入后调整，这里为了精简保证功能。
    
    # R0：姓名
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "{{ user_name }}"
    table.cell(0, 2).text = "工作单位及\n职务（职级）"
    table.cell(0, 3).text = "{{ dept_name }}"

    # R1：分管工作
    table.cell(1, 0).text = "从事或\n分管工作"
    table.cell(1, 1).merge(table.cell(1, 3))
    table.cell(1, 1).text = "各类检验检测与日常站内分管与相关行政事务。"

    # R2：大文本内容区域！
    table.cell(2, 0).text = "考\n\n\n核\n\n\n小\n\n\n结"
    table.cell(2, 1).merge(table.cell(2, 3))
    
    # 填充大内容并赋予仿宋和 {{ content }} 标志，保持排版一致性
    content_cell = table.cell(2, 1)
    content_cell.text = ""
    cp = content_cell.paragraphs[0]
    cp.paragraph_format.first_line_indent = Pt(32)
    cp.paragraph_format.line_spacing = Pt(28)
    cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    crun = cp.add_run("{{ content }}")
    crun.font.name = "仿宋_GB2312"
    crun.font.size = Pt(14)
    rpr = crun._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), '仿宋_GB2312'); rpr.insert(0, rf)

    # 底部意见栏
    sign_text = ["部门负责人\n评价意见", "主管领导\n意见", "院务会\n审定意见", "本人\n意见"]
    for i in range(4):
        r_idx = i + 3
        table.cell(r_idx, 0).text = sign_text[i]
        c2 = table.cell(r_idx, 1)
        c2.merge(table.cell(r_idx, 3))
        p = c2.add_paragraph("\n\n\n签名：                    年    月    日")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 遍历所有其余的表头，设定为黑体和居中
    for row in table.rows:
        for cell in row.cells:
            # 只对简短标题性质的格加粗和黑体化
            if len(cell.text) < 20 and "{{" not in cell.text:
                for p in cell.paragraphs:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = "黑体"
                        run.font.size = Pt(14)
                        run.bold = True
                        rpr = run._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), '黑体'); rpr.insert(0, rf)

    doc.save("sys_table_template.docx")

if __name__ == "__main__":
    create_table_template()
