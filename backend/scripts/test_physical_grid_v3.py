
import cv2
import numpy as np
from rapidocr_onnxruntime import RapidOCR
import fitz
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _apply_compact_style(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    
    run = p.add_run(str(text))
    run.font.name = "仿宋_GB2312"
    run.font.size = Pt(10)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), "仿宋_GB2312")
    rPr.append(rFonts)

def run_physical_test():
    pdf_path = r"F:\work-journal\tmp_SCAN_PAGE.pdf"
    doc_pdf = fitz.open(pdf_path)
    
    word_doc = Document()
    # A4 页面设置
    for section in word_doc.sections:
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
        section.top_margin = section.bottom_margin = Cm(1.5)
        section.left_margin = section.right_margin = Cm(2.0)

    ocr_engine = RapidOCR()

    for page_num in [0, 1]:  # 测试前两页
        page = doc_pdf[page_num]
        print(f"--- 处理第 {page_num + 1} 页 ---")
        # 提高到 3.0 倍缩放，确保小数字 1-9 不丢失
        pix = page.get_pixmap(matrix=fitz.Matrix(3.3, 3.3))
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
        img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
        img_width = img.shape[1]

        res, _ = ocr_engine(img)
        if not res: continue

        # 1. 物理行聚类 (Y 轴对齐)
        items = []
        for line in res:
            bbox, txt, _ = line
            y1 = min(p[1] for p in bbox)
            y2 = max(p[1] for p in bbox)
            cx = sum(p[0] for p in bbox) / 4
            cy = (y1 + y2) / 2
            items.append({'txt': txt.strip(), 'cx': cx, 'cy': cy, 'h': (y2-y1)})
        
        items.sort(key=lambda x: x['cy'])
        
        physical_rows = []
        if items:
            cur_row = [items[0]]
            for i in range(1, len(items)):
                # 若 Y 间距小于当前行平均高度的 60%，视为同一行
                if items[i]['cy'] - cur_row[-1]['cy'] < cur_row[-1]['h'] * 0.7:
                    cur_row.append(items[i])
                else:
                    physical_rows.append(cur_row)
                    cur_row = [items[i]]
            physical_rows.append(cur_row)

        print(f"识别到 {len(physical_rows)} 个物理行")

        # 2. 映射到 3 列网格
        # 探测结果 (3.3倍缩放后宽度约为 1965):
        # 序号 Cx 约 260
        # 名称 Cx 约 420
        # 价格 Cx 约 1040
        # 划分点设置为 330 和 850 (按比例换算自 1191 宽度的探测值)
        
        # 动态计算划分点 (基于当前图像宽度)
        p1 = img_width * (200 / 1191)  # 约 17% 处
        p2 = img_width * (550 / 1191)  # 约 46% 处
        
        table = word_doc.add_table(rows=len(physical_rows), cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        
        # 强制设置列宽比例 (1.5 : 6 : 2.5)
        # 总宽约 17.0cm (21.0 - 2.0*2)
        total_w = Cm(17.0)
        table.columns[0].width = int(total_w * 0.1)
        table.columns[1].width = int(total_w * 0.6)
        table.columns[2].width = int(total_w * 0.3)

        for r_idx, row_items in enumerate(physical_rows):
            # 将该物理行内的所有项按 X 轴分类
            cols_content = ["", "", ""]
            for it in row_items:
                if it['cx'] < p1:
                    cols_content[0] += it['txt'] + " "
                elif it['cx'] < p2:
                    cols_content[1] += it['txt'] + " "
                else:
                    cols_content[2] += it['txt'] + " "
            
            for c_idx in range(3):
                _apply_compact_style(table.cell(r_idx, c_idx), cols_content[c_idx].strip())
        
        # 简单分页处理
        if page_num == 0:
            word_doc.add_page_break()

    output_path = r"F:\work-journal\test_physical_output_v3.docx"
    word_doc.save(output_path)
    print(f"已生成测试文件: {output_path}")

if __name__ == "__main__":
    run_physical_test()
