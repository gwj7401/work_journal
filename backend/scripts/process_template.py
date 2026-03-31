import os
import docx
import traceback

def process_template():
    doc_path = r'f:\work-journal\docs\模板\郭卫江一季度-事业单位工作人员平时考核登记表.docx'
    out_path = r'f:\work-journal\docs\模板\月度总结模板.docx'
    
    if not os.path.exists(doc_path):
        print(f"找不到文件: {doc_path}")
        return

    try:
        doc = docx.Document(doc_path)

        # 全局替换“一季度”、“季度”到“月度”
        for p in doc.paragraphs:
            if '一季度' in p.text:
                p.text = p.text.replace('一季度', '月度')
            elif '季度' in p.text:
                p.text = p.text.replace('季度', '月度')

        longest_cell = None
        max_len = 0
        
        # 遍历替换表格中的文本，并找到最长的单元格（必定是个人工作总结输入区）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 替换可能包含的文字
                    if '一季度' in cell.text:
                        cell.text = cell.text.replace('一季度', '月度')
                    elif '季度' in cell.text:
                        cell.text = cell.text.replace('季度', '月度')

                    # 检查是否为大段文本（通常原文件中有大段总结）
                    length = len(cell.text)
                    if length > max_len:
                        max_len = length
                        longest_cell = cell

        # 如果找到了明显的长文本格子，或者只要存在单元格
        if longest_cell and max_len > 30:
            print(f"定位到原总结内容（前50个字符）: {longest_cell.text[:50]}...")
            
            # 清空它并放置 jinja2 变量 {{ content }}
            # 需要保留原有的段落样式
            if len(longest_cell.paragraphs) > 0:
                p = longest_cell.paragraphs[0]
                p.clear()
                run = p.add_run('{{ content }}')
                # 尽量匹配可能存在的字体，不过由于它是简单替换，格式应该会被docxtpl继承
                # 删除其他的段落
                for paragraph in longest_cell.paragraphs[1:]:
                    p_element = paragraph._element
                    p_element.getparent().remove(p_element)
                    paragraph._p = paragraph._element = None
            else:
                longest_cell.text = '{{ content }}'
            
            print("成功注入模板标签 {{ content }}。")
        else:
            print("警告：未找到明显的大段文字区域，可能定位失败。")

        doc.save(out_path)
        print(f"新模板生成完毕: {out_path}")

    except Exception as e:
        print("处理出错:")
        traceback.print_exc()

if __name__ == "__main__":
    process_template()
