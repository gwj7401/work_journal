"""
pdf2word_ocr.py — 跨平台 PDF→Word 物理布局还原引擎 (V3 RapidTable 集成版)

核心能力：
1. 跨平台兼容：兼容 Windows 和 Linux 的字体渲染。
2. 结构化表格还原（主路径）：使用 RapidTable 引擎直接输出 HTML → 解析 colspan/rowspan → 精准还原 Word 表格。
3. 坐标聚类回退：当 RapidTable 失败时，降级到基于文字坐标的虚拟网格算法。
4. 引擎适配：支持 RapidLayout + RapidTable + RapidOCR 协同工作。
"""

import sys
import os
import re
import platform
import fitz  # PyMuPDF
import numpy as np
from PIL import Image
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from rapidocr_onnxruntime import RapidOCR
from rapid_layout import RapidLayout
from rapid_table import RapidTable

# ──────────────────── 环境与配置 ────────────────────

IS_WINDOWS = (platform.system() == "Windows")

# 字体对照表（Windows 与 Linux 兼容方案）
FONT_MAP = {
    "title": "方正小标宋简体" if IS_WINDOWS else "Source Han Serif SC", # 标题：标宋或思源宋体
    "body": "仿宋_GB2312" if IS_WINDOWS else "FandolFang-Regular",        # 正文：仿宋
    "table": "仿宋_GB2312" if IS_WINDOWS else "Source Han Sans SC",        # 表格：仿宋或思源黑体
    "bold": "黑体" if IS_WINDOWS else "Source Han Sans SC"                # 粗体
}

# ──────────────────── 初始化引擎 (单例懒加载) ────────────────────

_engines = {'ocr': None, 'layout': None, 'table': None}

def _get_engine(type_name):
    global _engines
    if _engines[type_name] is None:
        try:
            if type_name == 'ocr': _engines['ocr'] = RapidOCR()
            elif type_name == 'layout': _engines['layout'] = RapidLayout()
            elif type_name == 'table': _engines['table'] = RapidTable()
        except Exception as e:
            raise Exception(f"{type_name} 引擎启动失败: {str(e)}")
    return _engines[type_name]

# ──────────────────── 通用样式工具 ────────────────────

def _apply_cell_style(cell, text, font_type="table", size=10):
    """为单元格应用紧凑样式：仿宋、10pt、单倍行距、无间距"""
    cell.text = ""  # 清空默认段落
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    run = p.add_run(str(text))
    font_name = FONT_MAP.get(font_type, "仿宋_GB2312")
    run.font.name = font_name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

# ──────────────────── 方案一：RapidTable 结构化 HTML 表格 ────────────────────

def _add_table_from_html(word_doc, html_str, ocr_results=None, total_width_px=800, total_height_px=600):
    """
    【RapidTable v2 结构还原路径】
    使用 RapidTable 识别的 HTML 结构，并结合全局 OCR 坐标进行补全注入。
    """
    if not html_str: return False

    soup = BeautifulSoup(html_str, 'html.parser')
    table_ele = soup.find('table')
    if not table_ele: return False

    rows = table_ele.find_all('tr')
    if not rows: return False

    # 1. 计算最大列数
    max_cols = 0
    row_cell_counts = []
    for row in rows:
        cells = row.find_all(['td', 'th'])
        row_len = 0
        for cell in cells:
            row_len += int(cell.get('colspan', 1))
        max_cols = max(max_cols, row_len)
        row_cell_counts.append(row_len)

    if max_cols == 0: return False

    # 2. 动态列宽权重计算 (复用之前的权重逻辑)
    col_weights = [1.0] * max_cols
    if ocr_results:
        x_bins = np.zeros(max_cols)
        bin_width = total_width_px / max_cols
        for line in ocr_results:
            bbox, text, _ = line
            cx = sum(p[0] for p in bbox) / 4
            idx = int(cx / bin_width)
            if 0 <= idx < max_cols:
                x_bins[idx] += len(text.strip())
        
        tw = sum(x_bins)
        if tw > 0:
            for i in range(max_cols):
                col_weights[i] = 0.4 + (x_bins[i] / tw) * (max_cols - 0.4 * max_cols)

    # 3. 创建 Word 表格
    num_rows = len(rows)
    doc_table = word_doc.add_table(rows=num_rows, cols=max_cols)
    doc_table.style = 'Table Grid'
    doc_table.autofit = False

    # 应用动态列宽
    word_total_width = Cm(15.6)
    weight_total = sum(col_weights)
    for i in range(max_cols):
        w = (col_weights[i] / weight_total) * word_total_width
        doc_table.columns[i].width = int(w)

    # 4. 样式与内容注入 (含 OCR 坐标映射补偿)
    occupancy = {}
    
    # 预处理 OCR 结果，方便快速查找
    # 重点关注序号列 (通常在第一列)
    ocr_items = []
    if ocr_results:
        for line in ocr_results:
            bbox, txt, _ = line
            x1, y1 = min(p[0] for p in bbox), min(p[1] for p in bbox)
            x2, y2 = max(p[0] for p in bbox), max(p[1] for p in bbox)
            ocr_items.append({'txt': txt.strip(), 'rect': (x1, y1, x2, y2), 'cx': (x1+x2)/2, 'cy': (y1+y2)/2})

    for r_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        col_ptr = 0

        for col_node in cells:
            while occupancy.get((r_idx, col_ptr), False):
                col_ptr += 1
            if col_ptr >= max_cols: break

            colspan = int(col_node.get('colspan', 1))
            rowspan = int(col_node.get('rowspan', 1))
            
            # 获取 HTML 中的文本
            html_text = col_node.get_text(separator="\n", strip=True)
            
            # --- 核心改进：OCR 坐标补偿 ---
            final_text = html_text
            if ocr_results and (not html_text or html_text == ""):
                # 估算该单元格在原图中的大概位置
                # y 分布按行等分，x 权重按 col_weights 分布
                y_start = (r_idx / num_rows) * total_height_px
                y_end = ((r_idx + rowspan) / num_rows) * total_height_px
                
                # 允许一定的垂直容差，优先匹配 Y 轴
                matched_ocr = []
                for item in ocr_items:
                    # 如果中心点在该行范围内
                    if y_start - 10 <= item['cy'] <= y_end + 10:
                        # 且在 X 轴的第一列区间内 (针对序号丢失最严重的场景)
                        if col_ptr == 0 and item['cx'] < (total_width_px / max_cols) * 0.8:
                            matched_ocr.append(item['txt'])
                
                if matched_ocr and not html_text:
                    final_text = "\n".join(matched_ocr)
                    # print(f"    [补全] 行{r_idx} 列{col_ptr} -> {final_text}")

            # 标记占位
            for r in range(r_idx, min(r_idx + rowspan, num_rows)):
                for c in range(col_ptr, min(col_ptr + colspan, max_cols)):
                    occupancy[(r, c)] = True

            try:
                cell = doc_table.cell(r_idx, col_ptr)
                if colspan > 1 or rowspan > 1:
                    cell.merge(doc_table.cell(min(r_idx + rowspan - 1, num_rows - 1), min(col_ptr + colspan - 1, max_cols - 1)))
                _apply_cell_style(cell, final_text)
            except Exception: pass

            col_ptr += colspan

    word_doc.add_paragraph().paragraph_format.space_before = Pt(6)
    print(f"  [RapidTable v2] ✅ 成功写入结构化表格 ({num_rows}行 × {max_cols}列)")
    return True

# ──────────────────── 方案二：坐标聚类回退 ────────────────────

def _add_table_by_physical_grid_v3(word_doc, ocr_results, total_width_px):
    """
    【Physical Grid V3 物理网格路径】
    使用 Y 轴物理行聚类 + X 轴固定 3 列映射。
    """
    if not ocr_results: return False

    # 1. 物理行聚类 (Y 轴对齐)
    items = []
    for line in ocr_results:
        bbox, txt, _ = line
        y1 = min(p[1] for p in bbox)
        y2 = max(p[1] for p in bbox)
        cx = sum(p[0] for p in bbox) / 4
        cy = (y1 + y2) / 2
        items.append({'txt': txt.strip(), 'cx': cx, 'cy': cy, 'h': (y2-y1)})
    
    items.sort(key=lambda x: x['cy'])
    
    physical_rows = []
    if items:
        cur_row = [items[0]]
        for i in range(1, len(items)):
            # 若 Y 间距小于当前行高度的 70%，视为同一行
            if items[i]['cy'] - cur_row[-1]['cy'] < cur_row[-1]['h'] * 0.7:
                cur_row.append(items[i])
            else:
                physical_rows.append(cur_row)
                cur_row = [items[i]]
        physical_rows.append(cur_row)

    # 2. 映射到 3 列网格 (17% 和 46% 处的划分点)
    p1 = total_width_px * 0.17
    p2 = total_width_px * 0.46
    
    table = word_doc.add_table(rows=len(physical_rows), cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    
    # 强制设置列宽比例 (1 : 6 : 3)
    total_w_cm = Cm(17.0)
    table.columns[0].width = int(total_w_cm * 0.1)
    table.columns[1].width = int(total_w_cm * 0.6)
    table.columns[2].width = int(total_w_cm * 0.3)

    for r_idx, row_items in enumerate(physical_rows):
        cols_content = ["", "", ""]
        for it in row_items:
            if it['cx'] < p1:
                cols_content[0] += it['txt'] + " "
            elif it['cx'] < p2:
                cols_content[1] += it['txt'] + " "
            else:
                cols_content[2] += it['txt'] + " "
        
        for c_idx in range(3):
            _apply_cell_style(table.cell(r_idx, c_idx), cols_content[c_idx].strip())
    
    word_doc.add_paragraph() # 间隙
    print(f"  [PhysicalGridV3] ✅ 物理网格处理完成 ({len(physical_rows)} 行)")
    return True

def _add_grid_table_by_coords(word_doc, ocr_results, total_width_px):
    """
    【坐标聚类回退方案】
    当 RapidTable 不可用时，基于 OCR 文字坐标进行 X/Y 轴聚类形成虚拟网格。
    """
    if not ocr_results: return

    # 1. 解析 OCR 坐标
    items = []
    for line in ocr_results:
        bbox, text, _ = line
        x1 = min(p[0] for p in bbox)
        x2 = max(p[0] for p in bbox)
        y1 = min(p[1] for p in bbox)
        y2 = max(p[1] for p in bbox)
        items.append({
            'text': text.strip(),
            'cx': (x1 + x2) / 2, 'cy': (y1 + y2) / 2,
            'x1': x1, 'x2': x2, 'y1': y1, 'y2': y2
        })

    # 2. Y 轴聚类分行
    items.sort(key=lambda x: x['cy'])
    rows = []
    if items:
        cur_row = [items[0]]
        for i in range(1, len(items)):
            if items[i]['cy'] - cur_row[-1]['cy'] < (items[i]['y2'] - items[i]['y1']) * 0.7:
                cur_row.append(items[i])
            else:
                rows.append(cur_row)
                cur_row = [items[i]]
        rows.append(cur_row)

    # 3. X 轴聚类分列
    all_cx = sorted([it['cx'] for it in items])
    cols_x = []
    if all_cx:
        cur_col_group = [all_cx[0]]
        for i in range(1, len(all_cx)):
            if all_cx[i] - sum(cur_col_group)/len(cur_col_group) < total_width_px * 0.05:
                cur_col_group.append(all_cx[i])
            else:
                cols_x.append(sum(cur_col_group)/len(cur_col_group))
                cur_col_group = [all_cx[i]]
        cols_x.append(sum(cur_col_group)/len(cur_col_group))

    num_cols = len(cols_x)
    num_rows = len(rows)
    if num_cols == 0 or num_rows == 0: return

    # 4. 创建表格并设置样式
    table = word_doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = False

    # 5. 动态列宽
    doc_width = Cm(15.6) 
    col_boundaries = []
    for i in range(len(cols_x) - 1):
        col_boundaries.append((cols_x[i] + cols_x[i+1]) / 2)
    
    widths_px = []
    prev_b = 0
    for b in col_boundaries:
        widths_px.append(b - prev_b)
        prev_b = b
    widths_px.append(total_width_px - prev_b)
    
    total_px = sum(widths_px)
    for i, w_px in enumerate(widths_px):
        if i < len(table.columns):
            table.columns[i].width = int(doc_width * (w_px / total_px))

    # 6. 填入数据
    for r_idx, row_items in enumerate(rows):
        for item in row_items:
            distances = [abs(item['cx'] - cx) for cx in cols_x]
            c_idx = distances.index(min(distances))
            cell = table.cell(r_idx, c_idx)
            
            existing = cell.text
            new_text = f"{existing}\n{item['text']}" if existing else item['text']
            _apply_cell_style(cell, new_text)

    word_doc.add_paragraph().paragraph_format.space_before = Pt(6)
    print(f"  [Fallback] 坐标聚类表格: {num_rows}行 × {num_cols}列")

# ──────────────────── 表格统一入口 ────────────────────

def _process_table_region(word_doc, img_region, region_width):
    """
    表格处理 v3：优先使用 Physical Grid V3 (物理坐标聚类)。
    """
    ocr_engine = _get_engine('ocr')
    ocr_res, _ = ocr_engine(img_region)
    region_height = img_region.shape[0]

    # V3 逻辑优先：针对扫描件特定 3 列结构
    try:
        if _add_table_by_physical_grid_v3(word_doc, ocr_res, region_width):
            return
    except Exception as e:
        print(f"  [Warn] Physical Grid V3 失败: {e}")

    # 回退到 RapidTable 或 旧版坐标聚类
    try:
        table_engine = _get_engine('table')
        table_result, _ = table_engine(img_region)
        html_str = None
        if isinstance(table_result, str): html_str = table_result
        elif hasattr(table_result, 'pred_htmls'): html_str = table_result.pred_htmls[0]
        elif isinstance(table_result, (list, tuple)): html_str = table_result[0]
            
        if html_str and _add_table_from_html(word_doc, html_str, ocr_results=ocr_res, total_width_px=region_width, total_height_px=region_height):
            return 
    except Exception as e:
        print(f"  [Warn] RapidTable 异常: {e}")
    
    _add_grid_table_by_coords(word_doc, ocr_res, region_width)

# ──────────────────── Word 文本处理 ────────────────────

def _add_normal_paragraph(word_doc, text, is_title=False):
    text = text.strip()
    if not text: return
    p = word_doc.add_paragraph()
    font_key = "title" if is_title else "body"
    font_name = FONT_MAP.get(font_key)
    
    if is_title:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing = Pt(28)
        if bool(re.search(r'[\u4e00-\u9fff]', text)) and not text.startswith(('-', '*', '•')):
            pf.first_line_indent = Pt(32)

    run = p.add_run(text)
    if is_title: run.bold = True
    run.font.name = font_name
    run.font.size = Pt(18 if is_title else 16)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

# ──────────────────── 主流程 ────────────────────

def pdf_to_word_ocr(pdf_path, output_docx_path, is_complex_mode=False):
    ocr_engine = _get_engine('ocr')
    layout_engine = _get_engine('layout')
    
    pdf_doc = fitz.open(pdf_path)
    word_doc = Document()

    try:
        # 设置页面
        for section in word_doc.sections:
            section.page_width, section.page_height = Cm(21.0), Cm(29.7)
            section.top_margin, section.bottom_margin = Cm(2.54), Cm(2.54)
            section.left_margin, section.right_margin = Cm(2.8), Cm(2.6)

        for page_num in range(len(pdf_doc)):
            print(f"⏳ 正在解析第 {page_num + 1} 页...")
            page = pdf_doc[page_num]
            # 提升缩放率至 3.3 倍，确保能识别扫描件中的极小文本 (如 1-9 序号)
            mat = fitz.Matrix(3.3, 3.3)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            img_array = np.array(img)
            # 确保传递给 OCR 的图像格式一致 (RapidOCR 内部处理 RGB/BGR，但 3.3x 需确保采样质量)

            # 1. 版面分析
            res_layout = layout_engine(img_array)
            layout_items = []
            if res_layout and hasattr(res_layout, 'boxes'):
                for box, cat in zip(res_layout.boxes, res_layout.class_names):
                    layout_items.append({'bbox': box, 'category': cat})
            
            # --- 布局容错增强：检查表格是否识别完整 ---
            need_complex_fallback = False
            if layout_items:
                all_y2 = [it['bbox'][3] for it in layout_items]
                max_layout_y = max(all_y2)
                if max_layout_y < img_array.shape[0] * 0.4:
                     temp_res, _ = ocr_engine(img_array)
                     if temp_res:
                         max_ocr_y = max([max(p[1] for p in line[0]) for line in temp_res])
                         if max_ocr_y > max_layout_y + 200:
                             print(f"  [Auto-Fix] 检测到布局识别不全 (Layout: {int(max_layout_y)} vs OCR: {int(max_ocr_y)})，切换全局模式...")
                             need_complex_fallback = True

            # 2. 逻辑分发
            if is_complex_mode or not layout_items or need_complex_fallback:
                print(f"  [Complex] 开启全局 RapidTable 解析...")
                # 全局模式也优先使用 RapidTable
                _process_table_region(word_doc, img_array, img_array.shape[1])
            else:
                # 按照 Y 轴坐标顺序处理版块
                layout_items.sort(key=lambda x: x['bbox'][1])
                for item in layout_items:
                    x1, y1, x2, y2 = map(int, item['bbox'])
                    region = img_array[y1:y2, x1:x2]

                    if item['category'] == 'table':
                        # 表格区域 → RapidTable 优先
                        _process_table_region(word_doc, region, x2 - x1)
                    else:
                        # 非表格区域 → OCR + 文本段落
                        res, _ = ocr_engine(region)
                        if not res: continue
                        full_text = " ".join([line[1] for line in res])
                        _add_normal_paragraph(word_doc, full_text, is_title=(item['category'] == 'title'))

            if page_num < len(pdf_doc) - 1:
                word_doc.add_page_break()

    finally:
        pdf_doc.close()

    word_doc.save(output_docx_path)
    print(f"🎉 转换完成: {output_docx_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        sys.exit(1)
    complex_mode = "--complex" in sys.argv
    pdf_to_word_ocr(sys.argv[1], sys.argv[2], is_complex_mode=complex_mode)
