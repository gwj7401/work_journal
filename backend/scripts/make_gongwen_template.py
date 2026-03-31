import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_gongwen_template():
    doc = Document()
    sec = doc.sections[0]
    
    # ==== 页面设置 ====
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.7)
    sec.bottom_margin = Cm(3.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(2.8)

    # ==== 标题 ====
    p_title = doc.add_paragraph()
    p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_title.paragraph_format.line_spacing = Pt(35)
    run_title = p_title.add_run("{{ year }}年{{ month }}月份工作总结")
    run_title.font.name = "方正小标宋简体"
    run_title.font.size = Pt(22) # 二号
    rpr = run_title._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), '方正小标宋简体'); rpr.insert(0, rf)

    # ==== 空行与署名 ====
    p_empty = doc.add_paragraph() # 空一行
    p_empty.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_empty.paragraph_format.line_spacing = Pt(28.8)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_sub.paragraph_format.line_spacing = Pt(28.8)
    run_sub = p_sub.add_run("{{ dept_name }}    {{ user_name }}")
    run_sub.font.name = "楷体_GB2312"
    run_sub.font.size = Pt(16) # 三号
    rpr_sub = run_sub._r.get_or_add_rPr(); rf_sub = OxmlElement('w:rFonts'); rf_sub.set(qn('w:eastAsia'), '楷体_GB2312'); rpr_sub.insert(0, rf_sub)
    
    # ==== 循环控制头 ====
    p_for = doc.add_paragraph()
    p_for.add_run("{%p for p in paras %}")

    # ==== 正文循环体 (每个自然段应用格式) ====
    p_content = doc.add_paragraph()
    # 标准公文：左起空二字（对于三号字，2个字符 = 32磅）
    p_content.paragraph_format.first_line_indent = Pt(32) 
    # 回行顶格即默认没有 hanging_indent 或 left_indent 配置
    p_content.paragraph_format.left_indent = 0
    p_content.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_content.paragraph_format.line_spacing = Pt(28.8)
    
    run_content = p_content.add_run("{{ p }}")
    run_content.font.name = "仿宋_GB2312"
    run_content.font.size = Pt(16) # 三号
    rpr_content = run_content._r.get_or_add_rPr(); rf_content = OxmlElement('w:rFonts'); rf_content.set(qn('w:eastAsia'), '仿宋_GB2312'); rpr_content.insert(0, rf_content)

    # ==== 循环控制尾 ====
    p_endfor = doc.add_paragraph()
    p_endfor.add_run("{%p endfor %}")

    out_path = r'f:\work-journal\docs\模板\月度总结模板.docx'
    if os.path.exists(out_path):
        os.remove(out_path)
    doc.save(out_path)
    print(f"生成的最新公文模板已保存为: {out_path}")

if __name__ == "__main__":
    create_gongwen_template()
