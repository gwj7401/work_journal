import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def convert_md_to_gongwen(md_path, docx_path):
    doc = Document()
    sec = doc.sections[0]
    
    # ==== 页面设置 ====
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3.7)
    sec.bottom_margin = Cm(3.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(32) # 首行缩进两字 (对于16磅字体)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28.8)
        
        if line.startswith('# '):
            p.paragraph_format.first_line_indent = 0
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = Pt(35)
            text = line[2:].strip()
            run = p.add_run(text)
            run.font.name = "方正小标宋简体"
            run.font.size = Pt(22) # 二号
            try:
                rpr = run._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), '方正小标宋简体'); rpr.insert(0, rf)
            except: pass
            
            # 添加副标题行或者留出空行
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p_sub.paragraph_format.line_spacing = Pt(28.8)
            run_sub = p_sub.add_run("郭卫江    2026年3月30日")
            run_sub.font.name = "楷体_GB2312"
            run_sub.font.size = Pt(16)
            try:
                rpr_sub = run_sub._r.get_or_add_rPr(); rf_sub = OxmlElement('w:rFonts'); rf_sub.set(qn('w:eastAsia'), '楷体_GB2312'); rpr_sub.insert(0, rf_sub)
            except: pass
            
        elif line.startswith('## '):
            text = line[3:].strip()
            # 解决如果有两级缩进，替换为公文编号等，这里直接加粗黑体即可
            run = p.add_run(text)
            run.font.name = "黑体"
            run.font.size = Pt(16) # 三号
            try:
                rpr = run._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), '黑体'); rpr.insert(0, rf)
            except: pass
            
        elif line.startswith('### '):
            text = line[4:].strip()
            run = p.add_run(text)
            run.font.name = "楷体_GB2312"
            run.bold = True
            run.font.size = Pt(16) # 三号
            try:
                rpr = run._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), '楷体_GB2312'); rpr.insert(0, rf)
            except: pass
            
        else:
            # 清理 markdown
            text = line.replace('**', '')
            
            # 兼容数字或符号列表，虽然首行缩进了，但是正文保持就可以
            run = p.add_run(text)
            run.font.name = "仿宋_GB2312"
            run.font.size = Pt(16) # 三号
            try:
                rpr = run._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), '仿宋_GB2312'); rpr.insert(0, rf)
            except: pass
        
    doc.save(docx_path)
    print(f"SUCCESS: {docx_path}")

if __name__ == "__main__":
    convert_md_to_gongwen(r"f:\work-journal\docs\2026-03-monthly-summary.md", r"f:\work-journal\docs\2026年3月份工作总结.docx")
