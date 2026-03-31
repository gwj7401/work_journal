from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import os

doc = Document()

# 设置整体中文字体
doc.styles['Normal'].font.name = u'仿宋'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
doc.styles['Normal'].font.size = Pt(16)  # 三号大概是16磅

# 标题
p_title = doc.add_paragraph()
p_title.alignment = 1 # 居中
run = p_title.add_run("国家机关政府部门公文格式标准")
run.font.name = u'方正小标宋简体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
run.font.size = Pt(22)  # 二号大概22磅

p_sub = doc.add_paragraph()
p_sub.alignment = 1
run_sub = p_sub.add_run("(2024 年最新版，建议收藏！)")
run_sub.font.size = Pt(14)

def add_para(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(32) #首行缩进两字
    p.paragraph_format.line_spacing_rule = 1 # 单倍
    r = p.add_run(text)

add_para("标题：字体为方正小标宋简体，字号为二号，由发文机关、发文事由、公文种类三部分组成。多行时要居中排布，做到词意完整，排列对称，长短适宜，间距恰当，标题排列应当使用梯形或菱形。(行距:固定值35磅)，题目名称下空一行写明单位和姓名(楷体GB_2312)。")
add_para("正文：内容要求准确地传达发文机关的有关方针、政策精神，写法力求简明扼要，条理清楚，实事求是，合乎文法，切忌冗长杂乱。文中的结构层次序数应准确掌握和使用。")
add_para("页面设置：上: 3.7  下: 3.5  左: 2.8  右: 2.6")
add_para("页眉：1.5     页脚：2.8")
add_para("每页 22 行，每行 28 字，行距：固定值 28.8 磅")
add_para("正文一般用 三号仿宋 GB-2312，段落及格式如下图所示：")
add_para("一级标题：黑体(三号)   “一、”")
add_para("二级标题：楷体 GB-2312 加粗   “（一）”")
add_para("三级标题：三号仿宋 GB-2312 加粗   “1.”，“(1)”")
add_para("公众号·万人公考资料库")

output_path = r"F:\Scan\_000045_大模型代拆原生提取版.docx"
doc.save(output_path)
print("DOCX_GENERATED:", output_path)
